import os
import json
import logging
from datetime import datetime

# Formatting tools
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

logger = logging.getLogger("etl_report_utils")

def ensure_dir(path):
    os.makedirs(os.path.dirname(path), exist_ok=True)

import re

def clean_xml_string(val):
    if val is None:
        return ""
    s = str(val)
    # Remove control characters (characters in range \x00-\x1f except \n, \t, \r)
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', s)
    return s

def clean_pdf_string(val, max_len=80):
    s = clean_xml_string(val)
    # Escape markup characters for ReportLab XML parser
    s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    if len(s) > max_len:
        s = s[:max_len-3] + "..."
    return s

def clean_docx_string(val, max_len=80):
    s = clean_xml_string(val)
    if len(s) > max_len:
        s = s[:max_len-3] + "..."
    return s

def clean_markdown_string(val, max_len=80):
    s = str(val) if val is not None else ""
    s = s.replace("|", "\\|")
    if len(s) > max_len:
        s = s[:max_len-3] + "..."
    return s

def generate_json_report(file_path: str, data: dict):
    ensure_dir(file_path)
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, default=str)
    logger.info(f"JSON report generated at {file_path}")

def generate_markdown_report(file_path: str, data: dict):
    ensure_dir(file_path)
    
    rca_text = ""
    for idx, rca in enumerate(data.get("root_cause_report", [])):
        rca_text += f"""
### Issue {idx+1}: {rca.get('issue')}
- **Root Cause**: {rca.get('root_cause')}
- **Business Impact**: {rca.get('business_impact')}
- **Technical Impact**: {rca.get('technical_impact')}
- **Recommendation**: {rca.get('recommendation')}
- **Confidence Score**: {rca.get('confidence')}%
"""

    md_content = f"""# ETL Pipeline Execution Report
**Batch ID**: `{data.get('batch_id')}` | **Date**: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}
**Status**: {data.get('pipeline_status')} | **Data Quality Score**: {data.get('quality_score')}%

---

## 1. Executive Summary
{data.get('business_summary') or 'No summary generated.'}

---

## 2. Dataset Profile & Initial Assessment
- **Dataset Name**: {data.get('dataset_name')}
- **Total Input Rows**: {data.get('metadata', {}).get('rows', 0)}
- **Total Columns**: {data.get('metadata', {}).get('columns', 0)}
- **Duplicate Rows Detected**: {data.get('duplicate_rows', 0)}
- **Missing Elements**: {sum(data.get('missing_values', {}).values()) if data.get('missing_values') else 0}

### Inferred Schema Layout
| Column Name | Inferred Datatype | Null Count |
|---|---|---|
"""
    
    schema_items = list(data.get('column_types', {}).items())
    for col, dtype in schema_items[:25]:
        nulls = data.get('missing_values', {}).get(col, 0)
        md_content += f"| {clean_markdown_string(col, 50)} | {clean_markdown_string(dtype, 30)} | {nulls} |\n"
    if len(schema_items) > 25:
        md_content += f"| ... and {len(schema_items) - 25} more columns ... | | |\n"

    md_content += f"""
---

## 3. Transformation Log
Total operations recorded: {len(data.get('transformation_history', []))}

| Target Column | Operation Applied | Purpose / Rationale |
|---|---|---|
"""
    transformations = data.get('transformation_history', [])
    for step in transformations[:25]:
        md_content += f"| {clean_markdown_string(step.get('column_name'), 40)} | {clean_markdown_string(step.get('new_value'), 40)} | {clean_markdown_string(step.get('reason'), 80)} |\n"
    if len(transformations) > 25:
        md_content += f"| ... and {len(transformations) - 25} more operations ... | | |\n"

    val_res = data.get('validation_results', {})
    md_content += f"""
---

## 4. Validation & Database Load
- **Records Successfully Loaded**: {val_res.get('rows_loaded', 0)}
- **Records Rejected**: {val_res.get('rows_rejected', 0)}
- **Staging DB Insert status**: {val_res.get('staging_status', 'N/A')}
- **Production DB Load status**: {val_res.get('production_status', 'N/A')}

### Executed SQL Operations Log
```sql
"""
    for sql in val_res.get('sql_logs', []):
        md_content += f"{sql}\n"
    md_content += "```\n"

    if val_res.get('rows_rejected', 0) > 0:
        md_content += f"""
### Rejected Records Sample (First 5 Rejects)
| Row Number | Reason for Rejection | Record Preview |
|---|---|---|
"""
        for rej in val_res.get('rejected_records', [])[:5]:
            md_content += f"| {rej.get('row_number')} | {rej.get('reason')} | `{json.dumps(rej.get('record'))}` |\n"

    if data.get('root_cause_report'):
        md_content += f"""
---

## 5. Root Cause Analysis (RCA)
{rca_text}
"""

    md_content += f"""
---

## 6. Business Insights & Recommendations
{data.get('business_insights') or 'No insights computed.'}

### Recommendations Checklist
"""
    for rec in (data.get('recommendations') or []):
        md_content += f"- [ ] {rec}\n"

    md_content += f"""
---

## 7. Performance Statistics
- **Total Job Duration**: {data.get('execution_time', 0.0):.2f} seconds
- **Database Engine**: MySQL Production
"""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    logger.info(f"Markdown report generated at {file_path}")


def generate_docx_report(file_path: str, data: dict):
    ensure_dir(file_path)
    doc = Document()
    
    # Title Section
    title = doc.add_paragraph()
    run = title.add_run("ETL Pipeline Execution & Intelligence Report")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    
    # Subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Batch: {data.get('batch_id')} | Status: {data.get('pipeline_status')} | Quality: {data.get('quality_score')}%")
    sub_run.italic = True
    
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(data.get('business_summary') or 'No summary generated.')
    
    doc.add_heading("2. Dataset Schema Assessment", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Dataset File: {data.get('dataset_name')}\n")
    p.add_run(f"Rows: {data.get('metadata', {}).get('rows', 0)} | Columns: {data.get('metadata', {}).get('columns', 0)}")
    
    # Table for Schema
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Column Name'
    hdr_cells[1].text = 'Datatype'
    hdr_cells[2].text = 'Null Elements'
    
    schema_items = list(data.get('column_types', {}).items())
    for col, dtype in schema_items[:25]:
        row_cells = table.add_row().cells
        row_cells[0].text = clean_docx_string(col, 50)
        row_cells[1].text = clean_docx_string(dtype, 30)
        row_cells[2].text = str(data.get('missing_values', {}).get(col, 0))
    if len(schema_items) > 25:
        row_cells = table.add_row().cells
        row_cells[0].text = f"... and {len(schema_items) - 25} more columns ..."
        row_cells[1].text = ""
        row_cells[2].text = ""
        
    doc.add_heading("3. Transformation Actions", level=1)
    table_t = doc.add_table(rows=1, cols=3)
    table_t.style = 'Light Shading Accent 1'
    hdr_t = table_t.rows[0].cells
    hdr_t[0].text = 'Column'
    hdr_t[1].text = 'Action'
    hdr_t[2].text = 'Rationale'
    
    transformations = data.get('transformation_history', [])
    for step in transformations[:25]:
        row_cells = table_t.add_row().cells
        row_cells[0].text = clean_docx_string(step.get('column_name'), 40)
        row_cells[1].text = clean_docx_string(step.get('new_value'), 40)
        row_cells[2].text = clean_docx_string(step.get('reason'), 80)
    if len(transformations) > 25:
        row_cells = table_t.add_row().cells
        row_cells[0].text = f"... and {len(transformations) - 25} more operations ..."
        row_cells[1].text = ""
        row_cells[2].text = ""
        
    doc.add_heading("4. Validation & Database Loading Logs", level=1)
    val_res = data.get('validation_results', {})
    p_val = doc.add_paragraph()
    p_val.add_run(f"Rows Loaded: {val_res.get('rows_loaded', 0)} | Rows Rejected: {val_res.get('rows_rejected', 0)}\n")
    p_val.add_run(f"Staging Status: {val_res.get('staging_status', 'N/A')} | Production Load: {val_res.get('production_status', 'N/A')}")
    
    # SQL logs
    doc.add_paragraph("Executed SQL statements:")
    for sql in val_res.get('sql_logs', []):
        doc.add_paragraph(sql, style='List Bullet')
        
    # Root Cause
    if data.get('root_cause_report'):
        doc.add_heading("5. Root Cause Analysis (RCA)", level=1)
        for rca in data.get('root_cause_report', []):
            p_rca = doc.add_paragraph()
            p_rca.add_run(f"Issue: {rca.get('issue')}\n").bold = True
            p_rca.add_run(f"Root Cause: {rca.get('root_cause')}\n")
            p_rca.add_run(f"Business Impact: {rca.get('business_impact')}\n")
            p_rca.add_run(f"Technical Impact: {rca.get('technical_impact')}\n")
            p_rca.add_run(f"Recommendation: {rca.get('recommendation')}\n")
            p_rca.add_run(f"Confidence Score: {rca.get('confidence')}%")

    doc.add_heading("6. Business Insights & Recommendations", level=1)
    doc.add_paragraph(data.get('business_insights') or 'No insights computed.')
    for rec in (data.get('recommendations') or []):
        doc.add_paragraph(rec, style='List Bullet')

    doc.save(file_path)
    logger.info(f"DOCX report generated at {file_path}")


def generate_pdf_report(file_path: str, data: dict):
    ensure_dir(file_path)
    
    # ReportLab Document setup
    doc = SimpleDocTemplate(file_path, pagesize=letter,
                            rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=22,
        leading=26,
        textColor=colors.HexColor('#1A365D'),
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        name='H1Style',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#2B6CB0'),
        spaceBefore=12,
        spaceAfter=8,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        name='BodyStyle',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#2D3748'),
        spaceAfter=6
    )

    bold_body_style = ParagraphStyle(
        name='BoldBodyStyle',
        parent=body_style,
        fontName='Helvetica-Bold'
    )
    
    story = []
    
    # Header block
    story.append(Paragraph("ETL Pipeline Execution & Intelligence Report", title_style))
    story.append(Paragraph(f"<b>Batch ID:</b> {data.get('batch_id')} | <b>Status:</b> {data.get('pipeline_status')} | <b>Quality Score:</b> {data.get('quality_score')}%", body_style))
    story.append(Spacer(1, 15))
    
    # Executive Summary
    story.append(Paragraph("1. Executive Summary", h1_style))
    story.append(Paragraph(data.get('business_summary') or 'No summary generated.', body_style))
    story.append(Spacer(1, 10))
    
    # Profile Info
    story.append(Paragraph("2. Dataset Assessment", h1_style))
    meta_info = f"<b>File name:</b> {data.get('dataset_name')} | <b>Total Rows:</b> {data.get('metadata', {}).get('rows', 0)} | <b>Total Columns:</b> {data.get('metadata', {}).get('columns', 0)}"
    story.append(Paragraph(meta_info, body_style))
    story.append(Spacer(1, 5))
    
    # Schema Table
    schema_data = [["Column Name", "Inferred Type", "Missing Elements"]]
    schema_items = list(data.get('column_types', {}).items())
    for col, dtype in schema_items[:25]:
        nulls = data.get('missing_values', {}).get(col, 0)
        schema_data.append([
            clean_pdf_string(col, 50),
            clean_pdf_string(dtype, 30),
            str(nulls)
        ])
    if len(schema_items) > 25:
        schema_data.append([
            f"... and {len(schema_items) - 25} more columns ...",
            "",
            ""
        ])
        
    t1 = Table(schema_data, colWidths=[2.5*inch, 2.5*inch, 2*inch])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#EDF2F7')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#2D3748')),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('FONTSIZE', (0,0), (-1,-1), 9),
    ]))
    story.append(t1)
    story.append(Spacer(1, 15))
    
    # Transformations
    story.append(Paragraph("3. Transformation Actions", h1_style))
    trans_data = [["Column", "Target Action", "Rationale / Meaning"]]
    transformations = data.get('transformation_history', [])
    for step in transformations[:25]:
        trans_data.append([
            clean_pdf_string(step.get('column_name'), 40),
            clean_pdf_string(step.get('new_value'), 40),
            clean_pdf_string(step.get('reason'), 80)
        ])
    if len(transformations) > 25:
        trans_data.append([
            f"... and {len(transformations) - 25} more operations ...",
            "",
            ""
        ])
    if len(trans_data) == 1:
        trans_data.append(["-", "No cleaning required", "Dataset clean"])
        
    t2 = Table(trans_data, colWidths=[2.0*inch, 2.0*inch, 3.0*inch])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#EDF2F7')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#2D3748')),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('FONTSIZE', (0,0), (-1,-1), 9),
    ]))
    story.append(t2)
    story.append(Spacer(1, 15))
    
    # PageBreak for RCA and details
    story.append(PageBreak())
    
    # 4. Storage Allocation & Formatting
    story.append(Paragraph("4. Storage Allocation & Formatting", h1_style))
    storage_info = f"• <b>Selected Storage Format:</b> {data.get('format_selected')}<br/>" \
                   f"• <b>Storage Justification:</b> {data.get('storage_reason')}<br/>" \
                   f"• <b>Physical Location Reference:</b> <font face='Courier'>{data.get('formatted_file_path')}</font>"
    story.append(Paragraph(storage_info, body_style))
    story.append(Spacer(1, 10))

    # 5. Validation & Database Load
    story.append(Paragraph("5. Validation & Database Load", h1_style))
    val_res = data.get('validation_results', {})
    val_summary = f"<b>Loaded Rows:</b> {val_res.get('rows_loaded', 0)} | <b>Rejected Rows:</b> {val_res.get('rows_rejected', 0)} | <b>Staging status:</b> {val_res.get('staging_status', 'N/A')} | <b>Production status:</b> {val_res.get('production_status', 'N/A')}"
    story.append(Paragraph(val_summary, body_style))
    
    # SQL logs paragraph
    sql_text = "<b>SQL Execution logs:</b><br/>"
    for sql in val_res.get('sql_logs', []):
        cleaned_sql = clean_pdf_string(sql, max_len=120)
        sql_text += f"- <font face='Courier'>{cleaned_sql}</font><br/>"
    story.append(Paragraph(sql_text, body_style))
    story.append(Spacer(1, 10))
    
    # Root Cause Analysis
    if data.get('root_cause_report'):
        story.append(Paragraph("6. Root Cause Analysis (RCA)", h1_style))
        for idx, rca in enumerate(data.get('root_cause_report', [])):
            story.append(Paragraph(f"<b>Issue {idx+1}: {rca.get('issue')}</b>", bold_body_style))
            story.append(Paragraph(f"• <b>Root Cause:</b> {rca.get('root_cause')}", body_style))
            story.append(Paragraph(f"• <b>Business Impact:</b> {rca.get('business_impact')}", body_style))
            story.append(Paragraph(f"• <b>Technical Impact:</b> {rca.get('technical_impact')}", body_style))
            story.append(Paragraph(f"• <b>Recommendation:</b> {rca.get('recommendation')}", body_style))
            story.append(Paragraph(f"• <b>AI Confidence Score:</b> {rca.get('confidence')}%", body_style))
            story.append(Spacer(1, 8))
            
    # Business Insights
    story.append(Paragraph("7. Business Insights & Recommendations", h1_style))
    insights = data.get('business_insights') or 'No insights computed.'
    if isinstance(insights, list):
        for ins in insights:
            story.append(Paragraph(f"• {ins}", body_style))
    else:
        story.append(Paragraph(str(insights), body_style))
        
    story.append(Spacer(1, 5))
    story.append(Paragraph("<b>Actionable Recommendations:</b>", bold_body_style))
    recommendations = data.get('recommendations') or []
    if isinstance(recommendations, list):
        for idx, rec in enumerate(recommendations):
            story.append(Paragraph(f"{idx+1}. {rec}", body_style))
    elif isinstance(recommendations, str):
        story.append(Paragraph(recommendations, body_style))
        
    # Performance Statistics
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"<b>Total Job Execution Duration:</b> {data.get('execution_time', 0.0):.2f} seconds", bold_body_style))
    
    doc.build(story)
    logger.info(f"PDF report generated at {file_path}")


def generate_txt_report(file_path: str, data: dict):
    ensure_dir(file_path)
    
    rca_text = ""
    for idx, rca in enumerate(data.get("root_cause_report", [])):
        rca_text += f"""
Issue {idx+1}: {rca.get('issue')}
- Root Cause: {rca.get('root_cause')}
- Business Impact: {rca.get('business_impact')}
- Technical Impact: {rca.get('technical_impact')}
- Recommendation: {rca.get('recommendation')}
- Confidence Score: {rca.get('confidence')}%
"""

    txt_content = f"""ETL Pipeline Execution Report
Batch ID: {data.get('batch_id')} | Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}
Status: {data.get('pipeline_status')} | Data Quality Score: {data.get('quality_score')}%

================================================================================

1. Executive Summary
{data.get('business_summary') or 'No summary generated.'}

================================================================================

2. Dataset Profile & Initial Assessment
- Dataset Name: {data.get('dataset_name')}
- Total Input Rows: {data.get('metadata', {}).get('rows', 0)}
- Total Columns: {data.get('metadata', {}).get('columns', 0)}
- Duplicate Rows Detected: {data.get('duplicate_rows', 0)}
- Missing Elements: {sum(data.get('missing_values', {}).values()) if data.get('missing_values') else 0}

Inferred Schema Layout:
"""
    
    schema_items = list(data.get('column_types', {}).items())
    for col, dtype in schema_items[:25]:
        nulls = data.get('missing_values', {}).get(col, 0)
        txt_content += f"  * {col}: Type={dtype}, Nulls={nulls}\n"
    if len(schema_items) > 25:
        txt_content += f"  * ... and {len(schema_items) - 25} more columns ...\n"

    txt_content += f"""
================================================================================

3. Transformation Log
Total operations recorded: {len(data.get('transformation_history', []))}
"""
    transformations = data.get('transformation_history', [])
    for step in transformations[:25]:
        txt_content += f"  * Column '{step.get('column_name')}': new_value='{step.get('new_value')}', reason='{step.get('reason')}'\n"
    if len(transformations) > 25:
        txt_content += f"  * ... and {len(transformations) - 25} more operations ...\n"

    val_res = data.get('validation_results', {})
    txt_content += f"""
================================================================================

4. Validation & Database Load
- Records Successfully Loaded: {val_res.get('rows_loaded', 0)}
- Records Rejected: {val_res.get('rows_rejected', 0)}
- Staging DB Insert status: {val_res.get('staging_status', 'N/A')}
- Production DB Load status: {val_res.get('production_status', 'N/A')}

Executed SQL Operations Log:
"""
    for sql in val_res.get('sql_logs', []):
        txt_content += f"  {sql}\n"

    if val_res.get('rows_rejected', 0) > 0:
        txt_content += "\nRejected Records Sample (First 5 Rejects):\n"
        for rej in val_res.get('rejected_records', [])[:5]:
            txt_content += f"  Row {rej.get('row_number')} | Reason: {rej.get('reason')} | Record: {json.dumps(rej.get('record'))}\n"

    if data.get('root_cause_report'):
        txt_content += f"""
================================================================================

5. Root Cause Analysis (RCA)
{rca_text}
"""

    txt_content += f"""
================================================================================

6. Business Insights & Recommendations
{data.get('business_insights') or 'No insights computed.'}

Recommendations Checklist:
"""
    for rec in (data.get('recommendations') or []):
        txt_content += f"- [ ] {rec}\n"

    txt_content += f"""
================================================================================

7. Performance Statistics
- Total Job Duration: {data.get('execution_time', 0.0):.2f} seconds
- Database Engine: MySQL Production
"""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(txt_content)
    logger.info(f"Text report generated at {file_path}")
