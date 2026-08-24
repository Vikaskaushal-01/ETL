import os
import logging
import uuid
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Header
from sqlalchemy.orm import Session
from backend.database.mysql import get_db
from backend.database.models import RagDocument
from backend.utils.account_utils import get_user_path
from typing import Optional

router = APIRouter(prefix="/rag", tags=["RAG Knowledge Base"])
logger = logging.getLogger("etl_rag")

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    # Try PyMuPDF (fitz)
    try:
        import fitz
        doc = fitz.open(file_path)
        for page in doc:
            t = page.get_text()
            if t:
                text += t + "\n"
        doc.close()
        if text.strip():
            logger.info("Extracted text successfully using PyMuPDF.")
            return text
    except Exception as e:
        logger.warning(f"PyMuPDF failed to extract text: {e}. Trying pypdf fallback...")
    
    # Try pypdf fallback
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        if text.strip():
            logger.info("Extracted text successfully using pypdf fallback.")
            return text
    except Exception as ex:
        logger.error(f"pypdf fallback failed: {ex}")
    
    return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        logger.info("Extracted text successfully from DOCX.")
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"docx extraction failed: {e}")
        return ""

def extract_text_from_plain(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            logger.info("Extracted plain text file.")
            return f.read()
    except Exception as e:
        logger.error(f"Plain text extraction failed: {e}")
        return ""

@router.post("/upload")
async def upload_rag_document(
    file: UploadFile = File(...), 
    db: Session = Depends(get_db), 
    x_user_email: Optional[str] = Header(None)
):
    """
    Upload a document (PDF, DOCX, TXT, MD, CSV, JSON) to RAG memory.
    Extracts text and saves the record in db.
    """
    email = x_user_email or "admin@controlai.net"
    sanitized = email.replace("@", "_").replace(".", "_")
    
    filename = file.filename
    _, ext = os.path.splitext(filename.lower())
    file_type = ext[1:] if ext else "unknown"
    
    allowed_types = ["pdf", "docx", "txt", "md", "csv", "json"]
    if file_type not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_type}. Supported types: {', '.join(allowed_types)}"
        )
    
    # Generate unique filename on disk to avoid conflicts
    unique_id = str(uuid.uuid4())[:8]
    base_name, _ = os.path.splitext(filename)
    safe_filename = f"{base_name}_{unique_id}{ext}"
    
    file_path = get_user_path(email, os.path.join("data", "rag_documents", safe_filename))
    db_file_path = f"Accounts/{sanitized}/data/rag_documents/{safe_filename}"
    
    try:
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
    except Exception as e:
        logger.error(f"Failed to save RAG file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save file on disk: {str(e)}")
    
    # Check file size limit (10MB)
    file_size = os.path.getsize(file_path)
    if file_size > 10 * 1024 * 1024:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit.")
    
    # Extract content
    extracted_content = ""
    if file_type == "pdf":
        extracted_content = extract_text_from_pdf(file_path)
    elif file_type == "docx":
        extracted_content = extract_text_from_docx(file_path)
    else:
        extracted_content = extract_text_from_plain(file_path)
        
    if not extracted_content.strip():
        # If extraction was empty, don't fail, but log warning
        logger.warning(f"No textual content could be extracted from {filename}")
        extracted_content = "[No text content found in document]"
        
    # Save database entry
    try:
        db_doc = RagDocument(
            filename=filename,
            file_type=file_type,
            file_path=db_file_path,
            content=extracted_content,
            uploaded_by=email
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)
    except Exception as db_err:
        if os.path.exists(file_path):
            os.remove(file_path)
        logger.error(f"Database registration for RAG doc failed: {db_err}")
        raise HTTPException(status_code=500, detail=f"Database registration failed: {str(db_err)}")
        
    return {
        "status": "Success",
        "id": db_doc.id,
        "filename": db_doc.filename,
        "file_type": db_doc.file_type,
        "upload_time": db_doc.upload_time
    }

@router.get("/documents")
def list_rag_documents(db: Session = Depends(get_db), x_user_email: Optional[str] = Header(None)):
    """
    List all documents in RAG knowledge base.
    """
    email = x_user_email or "admin@controlai.net"
    try:
        docs = db.query(RagDocument).filter(RagDocument.uploaded_by == email).order_by(RagDocument.upload_time.desc()).all()
        return [
            {
                "id": d.id,
                "filename": d.filename,
                "file_type": d.file_type,
                "file_path": d.file_path,
                "upload_time": d.upload_time,
                "word_count": len(d.content.split()) if d.content else 0
            }
            for d in docs
        ]
    except Exception as e:
        logger.error(f"Failed to list RAG documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/documents/{doc_id}")
def delete_rag_document(doc_id: int, db: Session = Depends(get_db), x_user_email: Optional[str] = Header(None)):
    """
    Delete a document from RAG knowledge base.
    """
    email = x_user_email or "admin@controlai.net"
    try:
        doc = db.query(RagDocument).filter(RagDocument.id == doc_id, RagDocument.uploaded_by == email).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete file from disk
        full_path = os.path.join(PROJECT_ROOT, doc.file_path)
        if os.path.exists(full_path):
            try:
                os.remove(full_path)
                logger.info(f"Deleted physical file from disk: {full_path}")
            except Exception as fe:
                logger.warning(f"Could not delete physical file: {fe}")
                
        # Delete db record
        db.delete(doc)
        db.commit()
        return {"status": "Success", "message": f"Document {doc_id} deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete RAG document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

from pydantic import BaseModel

class RagUrlRequest(BaseModel):
    url: str

@router.post("/upload/url")
async def upload_rag_url(
    req: RagUrlRequest,
    db: Session = Depends(get_db),
    x_user_email: Optional[str] = Header(None)
):
    """
    Fetch content from a URL and store it as text in the RAG database.
    """
    email = x_user_email or "admin@controlai.net"
    sanitized = email.replace("@", "_").replace(".", "_")
    url = str(req.url)
    
    import httpx
    from bs4 import BeautifulSoup
    
    # Fetch content
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(url, follow_redirects=True, timeout=15.0)
            if response.status_code != 200:
                raise HTTPException(status_code=400, detail=f"Failed to fetch RAG URL. Status code: {response.status_code}")
            html_content = response.text
    except Exception as e:
         raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {str(e)}")
         
    # Parse text from HTML
    try:
        # Use simple BeautifulSoup if bs4 is installed, otherwise standard regex strip
        try:
            soup = BeautifulSoup(html_content, "lxml")
            # remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            text_content = soup.get_text(separator="\n")
        except Exception:
            try:
                soup = BeautifulSoup(html_content, "html.parser")
                for script in soup(["script", "style"]):
                    script.decompose()
                text_content = soup.get_text(separator="\n")
            except Exception:
                # Regex fallback
                import re
                text_content = re.sub(r'<[^>]+>', '\n', html_content)
                
        # Clean up whitespace
        lines = (line.strip() for line in text_content.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text_content = "\n".join(chunk for chunk in chunks if chunk)
    except Exception as pe:
        logger.error(f"Failed to parse text from HTML: {pe}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text content: {str(pe)}")

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="No readable text content could be extracted from the URL.")
        
    # Limit content size (max 200KB of text)
    if len(text_content) > 200 * 1024:
        text_content = text_content[:200 * 1024] + "\n... [Content truncated due to size limit]"
        
    # Save database entry
    try:
        # Save a virtual file reference
        from urllib.parse import urlparse
        parsed = urlparse(url)
        filename = os.path.basename(parsed.path) or "webpage_link"
        if not filename or "." not in filename:
            filename = f"link_{uuid.uuid4().hex[:8]}.txt"
            
        db_doc = RagDocument(
            filename=f"URL: {url[:60]}...",
            file_type="url",
            file_path=url,
            content=text_content,
            uploaded_by=email
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)
    except Exception as db_err:
        logger.error(f"Database registration for RAG url failed: {db_err}")
        raise HTTPException(status_code=500, detail=f"Database registration failed: {str(db_err)}")
        
    return {
        "status": "Success",
        "id": db_doc.id,
        "filename": db_doc.filename,
        "file_type": db_doc.file_type,
        "upload_time": db_doc.upload_time
    }

