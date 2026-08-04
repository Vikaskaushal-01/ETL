import os
import json
import re
import logging
import time
import sqlite3
import pandas as pd
from sqlalchemy import text
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from backend.database.mysql import engine, SessionLocal
from backend.database.repository import log_agent_decision
from backend.core.llm import query_llm
from backend.utils.file_utils import read_dataset

logger = logging.getLogger("etl_storage_agent")

class StorageAgent:
    def __init__(self):
        self.role = "Senior Database & Storage Architect"
        self.name = "Intelligent Storage Agent"

    def run(self, clean_file_path: str, batch_id: str, metadata: dict = None) -> dict:
        start_time = time.time()
        logger.info(f"Running Intelligent Storage Agent on {clean_file_path} for batch {batch_id}")
        
        # 1. Read clean dataset
        df = read_dataset(clean_file_path)
        cols = list(df.columns)
        
        # 2. Analyze characteristics and select format
        # Heuristics:
        # - Relational entity datasets (e.g. containing customer_id or order_id) -> SQL
        # - Tabular datasets containing numerical/metric properties (e.g., quantities, prices, amounts) -> CSV
        # - Textual or unstructured datasets -> Word
        
        preview_data = df.head(3).to_dict(orient='records')
        col_types = {col: str(df[col].dtype) for col in cols}
        
        prompt = f"""
        You are a Senior Storage Architect. Analyze the characteristics of this clean dataset:
        Columns: {cols}
        Column Data Types: {col_types}
        Preview: {json.dumps(preview_data, default=str)}
        
        Select the optimal storage format among:
        - "CSV": Best for mostly numerical, statistical, or ledger tabular information.
        - "Word": Best for primarily textual, descriptive, or document-oriented datasets.
        - "SQL": Best for highly structured, relational databases requiring query scalability (e.g. master customer data, order lists).
        
        Return a JSON response with:
        1. format_selected: 'CSV', 'Word', or 'SQL'
        2. reason: clear technical justification for selecting this format.
        3. summary: a brief summary of the dataset.
        
        Return ONLY valid JSON.
        """
        
        system_instruction = "You are the Intelligent Storage Agent. Analyze dataset properties and determine optimal storage format as valid JSON."
        
        # Default fallback decisions
        if "customer_id" in cols or "order_id" in cols or "sale_id" in cols:
            # Let's check if it has lots of numerical values
            if "total_price" in cols or "quantity" in cols:
                default_format = "CSV"
                default_reason = "Dataset contains numerical transaction ledger information (quantity, unit_price, total_price) suited for flat CSV reports."
            else:
                default_format = "SQL"
                default_reason = "Dataset contains highly structured relational database keys (customer_id, order_id) mapping to specific entities."
        elif "text" in str(cols).lower() or "description" in str(cols).lower():
            default_format = "Word"
            default_reason = "Dataset is primarily textual or document-oriented."
        else:
            default_format = "CSV"
            default_reason = "Tabular data containing statistics and tabular attributes suited for spreadsheet-based CSV exports."
            
        try:
            llm_response = query_llm(prompt, system_instruction, json_mode=True)
            if "```json" in llm_response:
                llm_response = llm_response.split("```json")[1].split("```")[0].strip()
            elif "```" in llm_response:
                llm_response = llm_response.split("```")[1].split("```")[0].strip()
            decision = json.loads(llm_response.strip())
            format_selected = decision.get("format_selected", default_format).upper()
            storage_reason = decision.get("reason", default_reason)
        except Exception as e:
            logger.error(f"Error calling LLM in StorageAgent: {e}")
            format_selected = default_format.upper()
            storage_reason = default_reason

        # Standardize format naming
        if format_selected not in ["CSV", "WORD", "SQL"]:
            format_selected = default_format.upper()

        logger.info(f"Selected format: {format_selected}. Reason: {storage_reason}")
        
        # 3. Detect table name and set reference (No physical file stored for formats)
        dataset_type = "dataset"
        if "customer_id" in cols and "customer_name" in cols:
            dataset_type = "customers"
        elif "order_id" in cols and "customer_id" in cols:
            dataset_type = "orders"
        elif "sale_id" in cols and "order_id" in cols:
            dataset_type = "sales"
        else:
            if "customer_id" in cols:
                dataset_type = "customers"
            elif "order_id" in cols:
                dataset_type = "orders"
            elif "sale_id" in cols:
                dataset_type = "sales"
            else:
                dataset_type = "dataset"

        formatted_file_path = self._save_formatted_file_to_clean_folder(clean_file_path, format_selected, dataset_type, df)
        logger.info(f"Storage agent saved formatted dataset ({format_selected}) to clean folder: {formatted_file_path}")

        db = SessionLocal()
        sql_logs = []
        rejected_records = []
        rows_loaded = 0
        rows_rejected = 0
        
        # Pull existing FKs
        existing_customers = set()
        existing_orders = set()
        try:
            if dataset_type == "orders":
                res = db.execute(text("SELECT customer_id FROM customers")).fetchall()
                existing_customers = {row[0] for row in res}
            elif dataset_type == "sales":
                res = db.execute(text("SELECT order_id FROM orders")).fetchall()
                existing_orders = {row[0] for row in res}
        except Exception as e:
            logger.warning(f"Failed to fetch relational primary keys, using blank: {e}")

        staging_records = []
        production_records = []
        seen_pks = set()

        # Auto-provision missing parent entity keys if needed so foreign key constraints pass seamlessly
        missing_customers_to_stub = set()
        missing_orders_to_stub = set()

        for idx, row in df.iterrows():
            row_num = idx + 1
            row_dict = {k: (None if pd.isna(v) else v) for k, v in row.to_dict().items()}
            is_valid = True
            reject_reason = ""
            
            if dataset_type == "customers":
                pk_val = str(row_dict.get("customer_id", "")).strip()
                if not pk_val or pk_val == "nan" or pk_val == "None":
                    row_dict["customer_id"] = f"CUST_AUTO_{idx+1}"
                    pk_val = row_dict["customer_id"]
                if pk_val in seen_pks:
                    row_dict["customer_id"] = f"{pk_val}_dup_{idx+1}"
                seen_pks.add(row_dict["customer_id"])
                    
            elif dataset_type == "orders":
                pk_val = str(row_dict.get("order_id", "")).strip()
                fk_val = str(row_dict.get("customer_id", "")).strip()
                if not pk_val or pk_val == "nan" or pk_val == "None":
                    row_dict["order_id"] = f"ORD_AUTO_{idx+1}"
                    pk_val = row_dict["order_id"]
                if pk_val in seen_pks:
                    row_dict["order_id"] = f"{pk_val}_dup_{idx+1}"
                seen_pks.add(row_dict["order_id"])

                if not fk_val or fk_val == "nan" or fk_val == "None":
                    row_dict["customer_id"] = "CUST_DEFAULT"
                    fk_val = "CUST_DEFAULT"

                if fk_val not in existing_customers:
                    missing_customers_to_stub.add(fk_val)
                    existing_customers.add(fk_val)

            elif dataset_type == "sales":
                pk_val = str(row_dict.get("sale_id", "")).strip()
                fk_val = str(row_dict.get("order_id", "")).strip()
                if not pk_val or pk_val == "nan" or pk_val == "None":
                    row_dict["sale_id"] = f"SALE_AUTO_{idx+1}"
                    pk_val = row_dict["sale_id"]
                if pk_val in seen_pks:
                    row_dict["sale_id"] = f"{pk_val}_dup_{idx+1}"
                seen_pks.add(row_dict["sale_id"])

                if not fk_val or fk_val == "nan" or fk_val == "None":
                    row_dict["order_id"] = "ORD_DEFAULT"
                    fk_val = "ORD_DEFAULT"

                if fk_val not in existing_orders:
                    missing_orders_to_stub.add(fk_val)
                    existing_orders.add(fk_val)
            else:
                # Generic dataset: rows are all valid
                is_valid = True

            status = "Valid" if is_valid else "Rejected"
            if not is_valid:
                rows_rejected += 1
                rejected_records.append({
                    "row_number": row_num,
                    "record": row_dict,
                    "reason": reject_reason
                })
            else:
                rows_loaded += 1
                production_records.append(row_dict)

            row_dict["batch_id"] = batch_id
            row_dict["row_number"] = row_num
            row_dict["validation_status"] = status
            staging_records.append(row_dict)

        # Database Loading Transaction
        try:
            # Delete staging
            db.execute(text(f"DELETE FROM staging_{dataset_type} WHERE batch_id = :b"), {"b": batch_id})
            sql_logs.append(f"DELETE FROM staging_{dataset_type} WHERE batch_id = '{batch_id}'")
            
            # Bulk staging inserts
            if staging_records:
                if dataset_type == "customers":
                    stmt = text("""
                        INSERT INTO staging_customers (customer_id, customer_name, email, phone, region, batch_id, `row_number`, validation_status)
                        VALUES (:customer_id, :customer_name, :email, :phone, :region, :batch_id, :row_number, :validation_status)
                    """)
                    db.execute(stmt, staging_records)
                elif dataset_type == "orders":
                    formatted_staging = [{**r, "order_date": str(r.get("order_date")), "total_amount": str(r.get("total_amount"))} for r in staging_records]
                    stmt = text("""
                        INSERT INTO staging_orders (order_id, customer_id, order_date, status, total_amount, batch_id, `row_number`, validation_status)
                        VALUES (:order_id, :customer_id, :order_date, :status, :total_amount, :batch_id, :row_number, :validation_status)
                    """)
                    db.execute(stmt, formatted_staging)
                elif dataset_type == "sales":
                    formatted_staging = [{**r, "quantity": str(r.get("quantity")), "unit_price": str(r.get("unit_price")), "total_price": str(r.get("total_price")), "sale_date": str(r.get("sale_date"))} for r in staging_records]
                    stmt = text("""
                        INSERT INTO staging_sales (sale_id, order_id, product_id, quantity, unit_price, total_price, sale_date, batch_id, `row_number`, validation_status)
                        VALUES (:sale_id, :order_id, :product_id, :quantity, :unit_price, :total_price, :sale_date, :batch_id, :row_number, :validation_status)
                    """)
                    db.execute(stmt, formatted_staging)
                else:
                    formatted_staging = [
                        {
                            "batch_id": batch_id,
                            "row_number": r.get("row_number", idx + 1),
                            "data_json": json.dumps({k: v for k, v in r.items() if k not in ["batch_id", "row_number", "validation_status"]}, default=str),
                            "validation_status": r.get("validation_status", "Valid")
                        }
                        for idx, r in enumerate(staging_records)
                    ]
                    stmt = text("""
                        INSERT INTO staging_dataset (batch_id, `row_number`, data_json, validation_status)
                        VALUES (:batch_id, :row_number, :data_json, :validation_status)
                    """)
                    db.execute(stmt, formatted_staging)
            
            sql_logs.append(f"INSERTED {len(staging_records)} records into staging_{dataset_type}.")

            # Load production valid records
            is_sqlite = "sqlite" in str(db.get_bind().url)

            # Auto-provision parent stub records if referenced foreign keys are missing in DB
            if missing_customers_to_stub:
                for mc in missing_customers_to_stub:
                    try:
                        if is_sqlite:
                            db.execute(text("INSERT OR IGNORE INTO customers (customer_id, customer_name, email, phone, region) VALUES (:c, 'Auto-Provisioned Customer', 'auto@customer.internal', 'N/A', 'Global')"), {"c": mc})
                        else:
                            db.execute(text("INSERT INTO customers (customer_id, customer_name, email, phone, region) VALUES (:c, 'Auto-Provisioned Customer', 'auto@customer.internal', 'N/A', 'Global') ON DUPLICATE KEY UPDATE customer_id=customer_id"), {"c": mc})
                    except Exception as stub_err:
                        logger.warning(f"Customer stub insert note: {stub_err}")
                db.commit()

            if missing_orders_to_stub:
                # Ensure a base customer stub exists for auto orders
                try:
                    if is_sqlite:
                        db.execute(text("INSERT OR IGNORE INTO customers (customer_id, customer_name, email, phone, region) VALUES ('CUST_DEFAULT', 'Auto-Provisioned Customer', 'auto@customer.internal', 'N/A', 'Global')"))
                    else:
                        db.execute(text("INSERT INTO customers (customer_id, customer_name, email, phone, region) VALUES ('CUST_DEFAULT', 'Auto-Provisioned Customer', 'auto@customer.internal', 'N/A', 'Global') ON DUPLICATE KEY UPDATE customer_id=customer_id"))
                    db.commit()
                except Exception as stub_err:
                    logger.warning(f"Base customer stub insert note: {stub_err}")

                for mo in missing_orders_to_stub:
                    try:
                        if is_sqlite:
                            db.execute(text("INSERT OR IGNORE INTO orders (order_id, customer_id, order_date, status, total_amount) VALUES (:o, 'CUST_DEFAULT', CURRENT_TIMESTAMP, 'Auto-Provisioned', 0.0)"), {"o": mo})
                        else:
                            db.execute(text("INSERT INTO orders (order_id, customer_id, order_date, status, total_amount) VALUES (:o, 'CUST_DEFAULT', CURRENT_TIMESTAMP, 'Auto-Provisioned', 0.0) ON DUPLICATE KEY UPDATE order_id=order_id"), {"o": mo})
                    except Exception as stub_err:
                        logger.warning(f"Order stub insert note: {stub_err}")
                db.commit()

            if production_records:
                if dataset_type == "customers":
                    if is_sqlite:
                        stmt = text("""
                            INSERT OR REPLACE INTO customers (customer_id, customer_name, email, phone, region)
                            VALUES (:customer_id, :customer_name, :email, :phone, :region)
                        """)
                    else:
                        stmt = text("""
                            INSERT INTO customers (customer_id, customer_name, email, phone, region)
                            VALUES (:customer_id, :customer_name, :email, :phone, :region)
                            ON DUPLICATE KEY UPDATE customer_name=VALUES(customer_name), email=VALUES(email), phone=VALUES(phone), region=VALUES(region)
                        """)
                    db.execute(stmt, production_records)
                elif dataset_type == "orders":
                    formatted_prod = []
                    for r in production_records:
                        try:
                            o_date = pd.to_datetime(r.get("order_date"))
                            if pd.notna(o_date):
                                o_date = o_date.to_pydatetime()
                            else:
                                o_date = None
                        except Exception:
                            o_date = None
                        formatted_prod.append({
                            **r,
                            "order_date": o_date,
                            "total_amount": float(r.get("total_amount")) if r.get("total_amount") else 0.0
                        })
                    if is_sqlite:
                        stmt = text("""
                            INSERT OR REPLACE INTO orders (order_id, customer_id, order_date, status, total_amount)
                            VALUES (:order_id, :customer_id, :order_date, :status, :total_amount)
                        """)
                    else:
                        stmt = text("""
                            INSERT INTO orders (order_id, customer_id, order_date, status, total_amount)
                            VALUES (:order_id, :customer_id, :order_date, :status, :total_amount)
                            ON DUPLICATE KEY UPDATE customer_id=VALUES(customer_id), order_date=VALUES(order_date), status=VALUES(status), total_amount=VALUES(total_amount)
                        """)
                    db.execute(stmt, formatted_prod)
                elif dataset_type == "sales":
                    formatted_prod = []
                    for r in production_records:
                        try:
                            s_date = pd.to_datetime(r.get("sale_date"))
                            if pd.notna(s_date):
                                s_date = s_date.to_pydatetime()
                            else:
                                s_date = None
                        except Exception:
                            s_date = None
                        formatted_prod.append({
                            **r,
                            "quantity": int(r.get("quantity")) if r.get("quantity") else 0,
                            "unit_price": float(r.get("unit_price")) if r.get("unit_price") else 0.0,
                            "total_price": float(r.get("total_price")) if r.get("total_price") else 0.0,
                            "sale_date": s_date
                        })
                    if is_sqlite:
                        stmt = text("""
                            INSERT OR REPLACE INTO sales (sale_id, order_id, product_id, quantity, unit_price, total_price, sale_date)
                            VALUES (:sale_id, :order_id, :product_id, :quantity, :unit_price, :total_price, :sale_date)
                        """)
                    else:
                        stmt = text("""
                            INSERT INTO sales (sale_id, order_id, product_id, quantity, unit_price, total_price, sale_date)
                            VALUES (:sale_id, :order_id, :product_id, :quantity, :unit_price, :total_price, :sale_date)
                            ON DUPLICATE KEY UPDATE order_id=VALUES(order_id), product_id=VALUES(product_id), quantity=VALUES(quantity), unit_price=VALUES(unit_price), total_price=VALUES(total_price), sale_date=VALUES(sale_date)
                        """)
                    db.execute(stmt, formatted_prod)
                else:
                    formatted_prod = [
                        {"business_columns": json.dumps(r, default=str)}
                        for r in production_records
                    ]
                    stmt = text("""
                        INSERT INTO production_dataset (business_columns)
                        VALUES (:business_columns)
                    """)
                    db.execute(stmt, formatted_prod)
            
            sql_logs.append(f"INSERTED/UPDATED {len(production_records)} valid records into production {dataset_type} table.")
            db.commit()
        except Exception as db_err:
            db.rollback()
            logger.error(f"Database load failed: {db_err}")
            sql_logs.append(f"TRANSACTION ROLLBACK due to: {str(db_err)}")
            rows_rejected = len(df)
            rows_loaded = 0
            rejected_records = [{"row_number": i+1, "record": df.iloc[i].to_dict(), "reason": f"DB Load Crash: {str(db_err)}"} for i in range(len(df))]
        finally:
            db.close()

        validation_status = "Success" if rows_rejected == 0 else "Passed with Warnings"
        if rows_loaded == 0:
            validation_status = "Failed"

        execution_time = time.time() - start_time
        
        # Log agent decision in DB
        db_log = SessionLocal()
        try:
            log_agent_decision(
                db_log,
                batch_id=batch_id,
                agent_name=self.name,
                task="Format dataset and synchronize DB",
                reasoning=f"Selected storage format {format_selected} based on content structure. Loaded {rows_loaded} rows successfully, rejected {rows_rejected} rows.",
                confidence=95.0,
                execution_time=execution_time
            )
        except Exception as e:
            logger.error(f"Failed logging agent decision: {e}")
        finally:
            db_log.close()

        logger.info(f"Storage agent processing complete. Format: {format_selected}. Saved at: {formatted_file_path}")

        return {
            "format_selected": format_selected,
            "formatted_file_path": formatted_file_path,
            "storage_reason": storage_reason,
            "storage_status": "Success" if formatted_file_path else "Failed",
            "rows_loaded": rows_loaded,
            "rows_rejected": rows_rejected,
            "validation_results": {
                "rows_loaded": rows_loaded,
                "rows_rejected": rows_rejected,
                "validation_status": validation_status,
                "sql_logs": sql_logs,
                "rejected_records": rejected_records,
                "staging_status": "Success" if rows_loaded > 0 or rows_rejected > 0 else "Failed",
                "production_status": "Success" if rows_loaded > 0 else "Failed",
                "dataset_type": dataset_type,
                "execution_time": execution_time
            }
        }

    def _save_formatted_file_to_clean_folder(self, clean_file_path: str, format_selected: str, dataset_type: str, df: pd.DataFrame) -> str:
        try:
            PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
            clean_dir = os.path.join(PROJECT_ROOT, "cleaned data")
            os.makedirs(clean_dir, exist_ok=True)
            
            base_name = os.path.basename(clean_file_path)
            base_no_ext, _ = os.path.splitext(base_name)
            
            # Always save/export a copy in Microsoft Word (.docx) format to the Cleaned Data folder
            try:
                word_export_path = os.path.join(clean_dir, f"{base_no_ext}.docx")
                docx_doc = Document()
                docx_doc.add_heading(f"Cleaned Dataset: {base_no_ext}", level=0)
                docx_doc.add_paragraph(f"Formatted and structured dataset generated by Intelligent Storage Agent ({format_selected} format).")
                
                # Create Table (limit to 100 rows for size / speed reasons)
                rows_count, cols_count = df.shape
                preview_limit = min(100, rows_count)
                table = docx_doc.add_table(rows=preview_limit + 1, cols=cols_count)
                table.style = 'Table Grid'
                
                # Add Header
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                    
                # Add Data Rows
                for r_idx in range(preview_limit):
                    row_cells = table.rows[r_idx + 1].cells
                    for c_idx in range(cols_count):
                        val = df.iloc[r_idx, c_idx]
                        row_cells[c_idx].text = "" if pd.isna(val) else str(val)
                        
                docx_doc.save(word_export_path)
                logger.info(f"Saved formatted Word document to: {word_export_path}")
            except Exception as e_docx:
                logger.error(f"Failed exporting docx backup copy: {e_docx}")

            if format_selected == "WORD":
                return os.path.join(clean_dir, f"{base_no_ext}.docx").replace("\\", "/")

            elif format_selected == "SQL":
                target_path = os.path.join(clean_dir, f"{base_no_ext}.sql")
                sql_statements = [
                    f"-- Structured SQL Script Export for {base_no_ext}\n",
                    f"-- Table: {dataset_type}\n\n"
                ]
                
                col_defs = []
                for col in df.columns:
                    dtype_str = str(df[col].dtype)
                    if "int" in dtype_str:
                        col_type = "INT"
                    elif "float" in dtype_str:
                        col_type = "DECIMAL(10, 2)"
                    elif "date" in dtype_str or "time" in dtype_str:
                        col_type = "DATETIME"
                    else:
                        col_type = "VARCHAR(255)"
                    col_defs.append(f"  `{col}` {col_type}")
                    
                create_stmt = f"CREATE TABLE IF NOT EXISTS `{dataset_type}` (\n" + ",\n".join(col_defs) + "\n);\n\n"
                sql_statements.append(create_stmt)
                
                for _, row in df.iterrows():
                    val_strs = []
                    for v in row:
                        if pd.isna(v) or v is None:
                            val_strs.append("NULL")
                        elif isinstance(v, (int, float)):
                            val_strs.append(str(v))
                        else:
                            safe_val = str(v).replace("'", "''")
                            val_strs.append(f"'{safe_val}'")
                    stmt = f"INSERT INTO `{dataset_type}` (`{'`, `'.join(df.columns)}`) VALUES ({', '.join(val_strs)});\n"
                    sql_statements.append(stmt)
                    
                with open(target_path, "w", encoding="utf-8") as sf:
                    sf.write("".join(sql_statements))
                logger.info(f"Saved formatted SQL script to: {target_path}")
                return target_path.replace("\\", "/")
 
            else: # Default CSV
                target_path = os.path.join(clean_dir, f"{base_no_ext}.csv")
                df.to_csv(target_path, index=False)
                logger.info(f"Saved formatted CSV dataset to: {target_path}")
                return target_path.replace("\\", "/")
 
        except Exception as e:
            logger.error(f"Failed generating formatted file in clean folder: {e}")
            return clean_file_path
